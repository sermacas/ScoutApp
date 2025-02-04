import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
from shiny import App, render, ui, reactive
from shiny.ui import HTML
from datetime import datetime 
import pandas as pd
from shiny import App, reactive, render, ui
from htmltools import HTML
import json
from datetime import datetime
from pytube import Search
import requests
from bs4 import BeautifulSoup
import nest_asyncio
from shiny import run_app
from docx import Document
import os

nest_asyncio.apply()

# Configuración de Google Sheets
SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

# Cargar la variable de entorno que contiene las credenciales de Google
creds_json = os.getenv("GOOGLE_APPLICATION_CREDENTIALS_JSON")

# Verifica si la variable fue cargada correctamente
if creds_json is None:
    print("Error: La variable de entorno no está configurada correctamente.")
else:
    print("La variable de entorno se ha cargado correctamente.")

# Convierte el JSON cargado en un diccionario y crea las credenciales
creds_dict = json.loads(creds_json)
# Crear las credenciales con el alcance adecuado
creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)

# Autorizar el cliente de gspread con las credenciales
client = gspread.authorize(creds)

# Autoriza el cliente de gspread con las credenciales
client = gspread.authorize(creds)
SHEET_ID = "1MXuIF81o1Ts_QbEdhm0X790p4qRYxxtgoi3ufRwzEK8"

# Obtener datos desde Google Sheets
workbook = client.open_by_key(SHEET_ID)
worksheet = workbook.sheet1
values = worksheet.get_all_values()
df_jugadores = pd.DataFrame(values[1:], columns=values[0])
INFORMES_SHEET_NAME = "Historial"

# 🔹 Función para cargar informes desde Google Sheets
def cargar_historial():
    workbook = client.open_by_key(SHEET_ID)
    worksheet = workbook.worksheet(INFORMES_SHEET_NAME)
    
    values = worksheet.get_all_values()
    if len(values) < 2:
        return {}

    df_informes = pd.DataFrame(values[1:], columns=values[0])
    historial = {}
    
    for _, row in df_informes.iterrows():
        jugador = row["Jugador"]
        informe = {"Fecha": row["Fecha"], "Título": row["Título"], "Texto": row["Texto"]}
        
        if jugador not in historial:
            historial[jugador] = []
        historial[jugador].append(informe)

    return historial

# 🔹 Función para guardar informes en Google Sheets
def guardar_historial(historial):
    workbook = client.open_by_key(SHEET_ID)
    
    try:
        worksheet = workbook.worksheet(INFORMES_SHEET_NAME)
    except gspread.exceptions.WorksheetNotFound:
        worksheet = workbook.add_worksheet(title=INFORMES_SHEET_NAME, rows="1000", cols="4")

    df_historial = pd.DataFrame([
        {"Jugador": jugador, **informe}
        for jugador, informes in historial.items()
        for informe in informes
    ])

    values_actualizados = [df_historial.columns.tolist()] + df_historial.values.tolist()

    worksheet.clear()
    worksheet.update(values_actualizados)

    print("Historial de informes actualizado en Google Sheets.")

# 🔹 Estado reactivo para historial
historial_data = reactive.Value(cargar_historial())

# 🔹 Función para guardar un informe
def guardar_informe():
    jugador = input.jugador_informe()
    titulo = input.titulo_informe()
    texto = input.texto_informe()

    if not jugador or not titulo or not texto:
        return  

    fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    nuevo_informe = {"Fecha": fecha, "Título": titulo, "Texto": texto}

    historial_actual = historial_data.get().copy()
    historial_actual.setdefault(jugador, []).append(nuevo_informe)

    guardar_historial(historial_actual)
    historial_data.set(historial_actual)

    ui.update_text("titulo_informe", value="")
    ui.update_text("texto_informe", value="")
    
# Calcular la edad
current_year = datetime.now().year
df_jugadores["year_of_birth"] = df_jugadores["Birthdate"].astype(str).str[:4]
df_jugadores["year_of_birth"] = pd.to_numeric(df_jugadores["year_of_birth"], errors='coerce')
df_jugadores["Edad"] = current_year - df_jugadores["year_of_birth"]

df_jugadores['Club Contract'] = pd.to_datetime(df_jugadores['Club Contract'], errors='coerce').dt.date

# Reemplazar NaN en la columna "Contacto" por "No"
df_jugadores["Contacto"] = df_jugadores["Contacto"].fillna("No")

# Definir colores para "Sí" y "No" en la columna "Contacto"
contacto_colors = {
    "Si": "background-color: #49bb25 ; color: white;",  # Verde para "Sí"
    "No": "background-color: #dc3545; color: white;"   # Rojo para "No"
}

# Convertir la columna 'Birthdate' a datetime
df_jugadores['Birthdate'] = pd.to_datetime(df_jugadores['Birthdate'], errors='coerce')

# Calcular la edad de los jugadores
def calcular_edad(fecha_nacimiento):
    if pd.isna(fecha_nacimiento):
        return None
    hoy = datetime.today().date()
    edad = hoy.year - fecha_nacimiento.year - ((hoy.month, hoy.day) < (fecha_nacimiento.month, fecha_nacimiento.day))
    return edad

df_jugadores['Edad'] = df_jugadores['Birthdate'].apply(calcular_edad)

# Obtener lista de nacionalidades únicas, dividiendo combinaciones y ordenando
nacionalidades = sorted(set(sum([str(n).split(',') for n in df_jugadores["Nationality"].dropna().unique()], [])))
df_jugadores['Año Fin Contrato'] = df_jugadores['Club Contract'].apply(lambda x: x.year if pd.notna(x) else None)

# Definir colores para cada valor en "Assessment"
assessment_colors = {
    "Seguir": "background-color: #BB252E; color: white;",
    "Llamar": "background-color: #17a2b8; color: white;",
    "NaN": "background-color: #ffc107; color: black;",
    "All Iron Sports": "background-color: #49BB25; color: white;"
}

# Convertir diccionario de colores en JSON para JavaScript
colors_json = json.dumps(assessment_colors)

# Filtrar y ordenar las opciones
jugadores_unicos = sorted(df_jugadores["full_name"].dropna().unique().tolist())
columnas_unicas = sorted(df_jugadores.columns.tolist())
posicion_1_unica = sorted(df_jugadores["position_1"].dropna().unique().tolist())
posicion_2_unica = sorted(df_jugadores["position_2"].dropna().unique().tolist())
equipos_unicos = sorted(df_jugadores["Team"].dropna().unique().tolist())
nacionalidades_unicas = sorted(set(sum([str(n).split(',') for n in df_jugadores["Nationality"].dropna().unique()], [])))
agencias_unicas = sorted(df_jugadores["Agency"].dropna().unique().tolist())
seguimiento_unico = sorted(df_jugadores["Assessment"].dropna().unique().tolist())
categoria_unica = sorted(df_jugadores["source_sheet"].dropna().unique().tolist())

# Función para buscar videos en YouTube
def buscar_videos_youtube(nombre_jugador, equipo):
    try:
        query = f"{nombre_jugador} {equipo}"
        search_results = Search(query).results[:3]  # Limitar a 3 resultados
        videos = [{"title": video.title, "url": video.watch_url} for video in search_results]
        return videos
    except Exception as e:
        print(f"Error buscando videos: {e}")
        return []

# Función para buscar equipo y jugador, devolviendo solo el primer enlace
def buscar_equipo_jugador(jugador, equipo):
    query = f"{jugador} {equipo} lapreferente"
    url = f"https://duckduckgo.com/html/?q={query}"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        enlace = soup.find("a", class_="result__a", href=True)
        return enlace["href"] if enlace else None
    else:
        print("Error al realizar la solicitud")
        return None


def es_doble_nacionalidad(nacionalidad):
    return "," in nacionalidad  # Si hay una coma, es doble nacionalidad

informes_ui = ui.nav_panel(
    "Informes de Partidos",
    ui.layout_sidebar(
        ui.sidebar(
            ui.input_selectize("jugador_informe", "Seleccionar Jugador", choices=jugadores_unicos, remove_button=True),
            ui.input_text("titulo_informe", "Título del Informe"),
            ui.input_text_area("texto_informe", "Texto del Informe", placeholder="Escribe tu informe aquí..."),
            ui.input_action_button("guardar_informe", "Guardar Informe"),
            ui.download_button("descargar_historial", "Descargar Historial")
        ),
        ui.panel_well(
            ui.output_ui("historial_informes")
        )
    )
)

app_ui = ui.page_navbar(
    ui.nav_panel(
        "Tabla Jugadores",
        ui.layout_sidebar(
            ui.sidebar(
                ui.input_selectize("jugador", "Seleccionar Jugador", choices=jugadores_unicos, remove_button=True),
                ui.input_selectize("columna", "Seleccionar Columna", choices=columnas_unicas, remove_button=True, selected="Comentarios"),
                ui.input_text("nuevo_valor", "Nuevo Valor"),
                ui.input_action_button("actualizar", "Actualizar"),
                ui.panel_well(
                    ui.input_selectize("filtro_categoria", "Seleccionar categoria", choices=categoria_unica, remove_button=True, selected=None),
                    ui.input_text("filtro_nombre", "Filtrar por Nombre"),
                    ui.input_text("filtro_comentarios", "Filtrar por Comentarios"),
                    ui.input_selectize("filtro_posicion1", "Filtrar por Posición 1", choices=[""] + posicion_1_unica, remove_button=True),
                    ui.input_selectize("filtro_posicion2", "Filtrar por Posición 2", choices=[""] + posicion_2_unica, remove_button=True),
                    ui.input_selectize("filtro_equipo", "Filtrar por Equipo", choices=[""] + equipos_unicos, remove_button=True),
                    ui.input_selectize("filtro_nacionalidad", "Filtrar por Nacionalidad", choices=[""] + nacionalidades_unicas, remove_button=True),
                    ui.input_checkbox("doble_nacionalidad", "Mostrar solo jugadores con doble nacionalidad", False),
                    ui.input_selectize("filtro_agencia", "Filtrar por Agencia", choices=[""] + agencias_unicas, remove_button=True),
                    ui.input_slider("filtro_edad_min", "Edad Mínima", min=12, max=24, value=14),
                    ui.input_slider("filtro_edad_max", "Edad Máxima", min=12, max=25, value=25),
                    ui.input_selectize("filtro_seguimiento", "Filtrar por Seguimiento", choices=[""] + seguimiento_unico, remove_button=True),
                    ui.input_action_button("restablecer_filtros", "Restablecer Filtros"),
                )
            ),
            ui.panel_well(
                ui.output_ui("tabla_actualizada")
            )
        )
    ),
    ui.nav_panel(
        "Detalle Jugador",
        ui.layout_sidebar(
            ui.sidebar(
                ui.input_selectize("jugador_detalle", "Seleccionar Jugador", choices=jugadores_unicos, remove_button=True),
            ),
            ui.panel_well(
                ui.output_ui("detalle_jugador"),
                ui.output_ui("videos_jugador"),
                ui.output_ui("lapreferente_jugador")
            )
        )
    ),
    informes_ui,  # Añadir la nueva página de informes
    title="Scouting App"
)

def server(input, output, session):
    df_reactivo = reactive.Value(df_jugadores.copy())

    @reactive.effect
    @reactive.event(input.restablecer_filtros)
    def restablecer_filtros():
        ui.update_text("filtro_nombre", value="")
        ui.update_text("filtro_comentarios", value="")
        ui.update_select("filtro_posicion1", selected="")
        ui.update_select("filtro_posicion2", selected="")
        ui.update_select("filtro_equipo", selected="")
        ui.update_select("filtro_nacionalidad", selected="")
        ui.update_select("filtro_agencia", selected="")
        ui.update_slider("filtro_edad_min", value=18)
        ui.update_slider("filtro_edad_max", value=25)
        ui.update_select("filtro_seguimiento", selected="")
        ui.update_select("filtro_categoria", selected="")

    df_reactivo = reactive.Value(df_jugadores)
    @reactive.effect
    @reactive.event(input.actualizar)
    def actualizar_datos():
        df = df_reactivo().copy()
        jugador = input.jugador()
        columna = input.columna()
        nuevo_valor = input.nuevo_valor()

        if jugador in df["full_name"].values and columna in df.columns:
            # Obtener el índice de la fila y la columna en Google Sheets
            col_index = df.columns.get_loc(columna) + 1  # Índice de columna en Google Sheets (1-based)
            row_index = df[df["full_name"] == jugador].index[0] + 2  # Fila en Google Sheets (1-based)

            try:
                worksheet.update_cell(row_index, col_index, nuevo_valor)  # Actualizar solo la celda necesaria
                df.loc[df["full_name"] == jugador, columna] = nuevo_valor  # Actualizar DataFrame reactivo
                df_reactivo.set(df)
            except Exception as e:
                print(f"Error actualizando Google Sheets: {e}")  # Capturar errores para debug

    @output
    @render.ui
    def tabla_actualizada():
        df = df_reactivo().copy()

        # Aplicar filtros
        if input.filtro_nombre():
            df = df[df["full_name"].str.contains(input.filtro_nombre(), case=False, na=False)]
        if input.filtro_comentarios():
            df = df[df["Comentarios"].str.contains(input.filtro_comentarios(), case=False, na=False)]
        if input.filtro_posicion1():
            df = df[df["position_1"] == input.filtro_posicion1()]
        if input.filtro_posicion2():
            df = df[df["position_2"] == input.filtro_posicion2()]
        if input.filtro_equipo():
            df = df[df["Team"] == input.filtro_equipo()]
        if input.filtro_nacionalidad():
            df = df[df["Nationality"].str.contains(input.filtro_nacionalidad(), na=False)]
        if "Agency" in df.columns and input.filtro_agencia():
            df = df[df["Agency"] == input.filtro_agencia()]
        if "Assessment" in df.columns and input.filtro_seguimiento():
            df = df[df["Assessment"] == input.filtro_seguimiento()]
        if "source_sheet" in df.columns and input.filtro_categoria():
            df = df[df["source_sheet"] == input.filtro_categoria()]
        if input.doble_nacionalidad():
            df = df[df["Nationality"].apply(es_doble_nacionalidad)]

        if "Edad" in df.columns:
            edad_min = input.filtro_edad_min()
            edad_max = input.filtro_edad_max()
            df = df[(df["Edad"] >= edad_min) & (df["Edad"] <= edad_max)]

        # Función para colorear celdas y aplicar negrita al nombre
        def colorize(col, value):
            if col == "full_name":
                return f'<td><strong>{value}</strong></td>'  # Nombre en negrita
            elif col == "Contacto" and value in contacto_colors:
                return f'<td style="{contacto_colors[value]}">{value}</td>'
            elif value in assessment_colors:
                return f'<td style="{assessment_colors.get(value, "")}">{value}</td>'
            else:
                return f'<td>{value}</td>'

        # Generar tabla HTML con estilos mejorados
        table_html = '''
        <style>
            table.dataframe {
                border-collapse: collapse;
                width: 100%;
                font-size: 12px;  /* Reducir el tamaño del texto */
                margin-bottom: 20px;
                max-height: 800px;  /* Altura máxima para que no ocupe mucho espacio */
                overflow-y: auto;  /* Hacer que la tabla tenga scroll si es necesario */
                display: block;
            }
            table.dataframe th, table.dataframe td {
                border: 1px solid #ddd;
                padding: 4px;  /* Reducir padding para que la tabla sea más compacta */
                text-align: left;
            }
            table.dataframe th {
                background-color: #f2f2f2;
                color: #333;
            }
            table.dataframe tr:nth-child(even) {
                background-color: #f9f9f9;
            }
            table.dataframe tr:hover {
                background-color: #f1f1f1;
            }
        </style>
        <table class="dataframe">
        '''
        table_html += "<thead><tr>" + "".join(f"<th>{col}</th>" for col in df.columns) + "</tr></thead>"
        table_html += "<tbody>"
        for _, row in df.iterrows():
            table_html += "<tr>" + "".join(colorize(col, value) for col, value in row.items()) + "</tr>"
        table_html += "</tbody></table>"

        return HTML(table_html)

    @output
    @render.ui
    def detalle_jugador():
        jugador = input.jugador_detalle()
        if jugador:
            df = df_reactivo().copy()
            jugador_data = df[df["full_name"] == jugador].iloc[0]
            
            # Definir las secciones de información
            info_personal = {
                "Nombre Completo": jugador_data["full_name"],
                "Fecha de Nacimiento": jugador_data["Birthdate"],
                "Edad": jugador_data["Edad"],
                "Nacionalidad": jugador_data["Nationality"]
            }
            
            info_deportiva = {
                "Posición 1": jugador_data["position_1"],
                "Posición 2": jugador_data["position_2"],
                "Equipo": jugador_data["Team"],
                "Contrato con el Club": jugador_data["Club Contract"],
                "Agencia": jugador_data["Agency"]
            }
            
            info_seguimiento = {
                "Assessment": jugador_data["Assessment"],
                "Comentarios": jugador_data["Comentarios"]
            }
            
            # Función para generar una tarjeta de información
            def generar_tarjeta(titulo, datos):
                contenido = f'<div class="card" style="margin-bottom: 20px; padding: 15px; border: 1px solid #ddd; border-radius: 5px; background-color: #f9f9f9;">'
                contenido += f'<h4 style="margin-top: 0;">{titulo}</h4>'
                for clave, valor in datos.items():
                    contenido += f'<p><strong>{clave}:</strong> {valor}</p>'
                contenido += '</div>'
                return contenido
            
            # Generar las tarjetas
            tarjeta_personal = generar_tarjeta("Información Personal", info_personal)
            tarjeta_deportiva = generar_tarjeta("Datos Deportivos", info_deportiva)
            tarjeta_seguimiento = generar_tarjeta("Seguimiento", info_seguimiento)
            
            # Combinar las tarjetas en un solo HTML
            detalles_html = f'''
            <div style="display: flex; flex-direction: column; gap: 20px;">
                {tarjeta_personal}
                {tarjeta_deportiva}
                {tarjeta_seguimiento}
            </div>
            '''
            
            return HTML(detalles_html)
        else:
            return HTML("Selecciona un jugador para ver sus detalles.")

    @output
    @render.ui
    def videos_jugador():
        jugador = input.jugador_detalle()
        if jugador:
            df = df_reactivo().copy()
            equipo = df[df["full_name"] == jugador]["Team"].iloc[0]
            videos = buscar_videos_youtube(jugador, equipo)
            if videos:
                video_html = "<h3>Videos Relacionados</h3>"
                for video in videos:
                    video_html += f'<p><a href="{video["url"]}" target="_blank">{video["title"]}</a></p>'
                return HTML(video_html)
            else:
                return HTML("No se encontraron videos para este jugador.")
        else:
            return HTML("")

    @output
    @render.ui
    def lapreferente_jugador():
        jugador = input.jugador_detalle()
        if jugador:
            df = df_reactivo().copy()
            equipo = df[df["full_name"] == jugador]["Team"].iloc[0] if "Team" in df.columns else ""
            
            # Usamos la función buscar_equipo_jugador para obtener solo el primer enlace
            link = buscar_equipo_jugador(jugador, equipo)

            if link:
                lapreferente_html = f'<h3>Primer resultado en LaPreferente</h3><p><a href="{link}" target="_blank">{link}</a></p>'
                return HTML(lapreferente_html)
            else:
                return HTML("No se encontró ningún resultado en LaPreferente para este jugador.")
        else:
            return HTML("")
        
    # Diccionario reactivo para almacenar informes por jugador
    historial_data = reactive.Value(cargar_historial())

    @reactive.effect
    @reactive.event(input.guardar_informe)
    def guardar_informe():
        jugador = input.jugador_informe()
        titulo = input.titulo_informe()
        texto = input.texto_informe()
        
        if not jugador or not titulo or not texto:
            return  

        fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")  

        nuevo_informe = {
            "Fecha": fecha,
            "Título": titulo,
            "Texto": texto
        }
        
        historial_actual = historial_data.get().copy()
        historial_actual.setdefault(jugador, []).append(nuevo_informe)
        
        # Guardar en el archivo JSON
        guardar_historial(historial_actual)

        # Actualizar historial reactivo
        historial_data.set(historial_actual)

        # 🔹 Limpiar los campos después de guardar
        ui.update_text("titulo_informe", value="")
        ui.update_text("texto_informe", value="")

    @output
    @render.ui
    def historial_informes():
        jugador_seleccionado = input.jugador_informe()

        if not jugador_seleccionado:
            return HTML("<p>Seleccione un jugador para ver su historial de informes.</p>")

        informes = historial_data.get().get(jugador_seleccionado, [])

        if not informes:
            return HTML("<p>No hay informes para este jugador.</p>")

        informes_html = f"<h3>Historial de Informes de {jugador_seleccionado}</h3>"
        for informe in informes:
            informes_html += f"""
            <div style="border: 1px solid #ddd; padding: 10px; margin-bottom: 10px; border-radius: 5px;">
                <p><strong>Fecha:</strong> {informe['Fecha']}</p>
                <p><strong>Título:</strong> {informe['Título']}</p>
                <p>{informe['Texto']}</p>
            </div>
            """
        return HTML(informes_html)

    @output
    @render.download
    def descargar_historial():
        jugador_seleccionado = input.jugador_informe()

        if not jugador_seleccionado:
            return None

        informes = historial_data.get().get(jugador_seleccionado, [])
        
        if not informes:
            return None

        doc = Document()
        doc.add_heading(f'Historial de Informes - {jugador_seleccionado}', level=1)

        for informe in informes:
            doc.add_paragraph(f"📅 Fecha: {informe['Fecha']}", style="Normal")
            doc.add_paragraph(f"📝 Título: {informe['Título']}", style="Normal")
            doc.add_paragraph(f"{informe['Texto']}", style="Normal")
            doc.add_paragraph("\n" + "-"*40 + "\n")

        file_path = f"historial_{jugador_seleccionado}.docx"
        doc.save(file_path)

        return file_path

app = App(app_ui, server)